from docx import Document
import sys

doc = Document(r"C:/Users/alekc/Downloads/LIMINAL (2).docx")
for p in doc.paragraphs:
    if p.text.strip():
        print(p.text)
print("\n--- TABLES ---")
for table in doc.tables:
    for row in table.rows:
        print([cell.text for cell in row.cells])
    print("---")
